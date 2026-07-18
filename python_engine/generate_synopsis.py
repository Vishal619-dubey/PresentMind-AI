import json,sys,re,os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, ListFlowable, ListItem, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
from reportlab.lib.units import inch

data=json.loads(sys.argv[1]);docx_out=sys.argv[2];pdf_out=sys.argv[3]

def words(s): return str(s or '').strip()
def heading(k): return re.sub(r'([A-Z])',r' \1',k).replace('_',' ').title()

def section_text(key, minimum=3):
    value=data.get(key,'')
    if isinstance(value,list): return [words(x) for x in value if words(x)]
    return words(value)

# ---------- DOCX ----------
doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.7);sec.bottom_margin=Inches(.7);sec.left_margin=Inches(.75);sec.right_margin=Inches(.75)
styles=doc.styles
styles['Normal'].font.name='Georgia';styles['Normal'].font.size=Pt(11)
styles['Normal'].paragraph_format.space_after=Pt(6);styles['Normal'].paragraph_format.line_spacing=1.15
styles['Heading 1'].font.name='Georgia';styles['Heading 1'].font.color.rgb=RGBColor(50,48,132);styles['Heading 1'].font.size=Pt(18)
styles['Heading 2'].font.name='Georgia';styles['Heading 2'].font.color.rgb=RGBColor(70,66,160);styles['Heading 2'].font.size=Pt(14)

p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(data.get('title','Project Synopsis'));r.bold=True;r.font.name='Georgia';r.font.size=Pt(24);r.font.color.rgb=RGBColor(40,42,110)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('PROJECT SYNOPSIS');r.bold=True;r.font.size=Pt(16)
doc.add_paragraph('')
t=doc.add_table(rows=6,cols=2);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.style='Table Grid'
meta=[('Submitted By',data.get('studentName','Vishal Dubey')),('Programme',data.get('programme','Master of Computer Applications (MCA)')),('University',data.get('university','Integral University, Lucknow')),('Academic Session',data.get('session','2026–2027')),('Technology Stack',data.get('technologyStack','React, Node.js, Express, MongoDB, Python')),('Generated On',data.get('generatedOn',''))]
for row,(a,b) in zip(t.rows,meta):
    row.cells[0].text=a;row.cells[1].text=words(b);row.cells[0].vertical_alignment=row.cells[1].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row.cells[0].paragraphs[0].runs[0].bold=True
for _ in range(4): doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run('Intelligent Presentation & Research Workspace').italic=True
doc.add_page_break()

doc.add_heading('Declaration',1)
doc.add_paragraph(f"I, {data.get('studentName','Vishal Dubey')}, declare that the project synopsis entitled “{data.get('title','')}” represents an original proposed work prepared for academic and development purposes. The proposed design, objectives, methodology and implementation plan have been described to the best of my understanding. Any external concepts, libraries or research ideas used during implementation will be properly acknowledged in the final report.")
doc.add_paragraph('\nStudent Signature: ____________________\nDate: ____________________')
doc.add_page_break()

doc.add_heading('Table of Contents',1)
order=['abstract','introduction','backgroundAndMotivation','problemStatement','existingSystem','limitationsOfExistingSystem','proposedSystem','objectives','scope','modules','functionalRequirements','nonFunctionalRequirements','methodology','systemArchitecture','databaseDesign','hardwareRequirements','softwareRequirements','testingStrategy','expectedResults','advantages','limitations','futureScope','workPlan','conclusion','references']
for i,k in enumerate(order,1): doc.add_paragraph(f'{i}. {heading(k)}',style='List Number')
doc.add_page_break()

for k in order:
    val=data.get(k)
    if not val: continue
    doc.add_heading(heading(k),1)
    if isinstance(val,list):
        for item in val:
            if isinstance(item,dict):
                title=item.get('title') or item.get('name') or item.get('phase') or 'Item'
                desc=item.get('description') or item.get('details') or item.get('duration') or ''
                p=doc.add_paragraph(style='List Bullet');p.add_run(str(title)).bold=True
                if desc:p.add_run(': '+str(desc))
            else: doc.add_paragraph(str(item),style='List Bullet')
    elif isinstance(val,dict):
        for sk,sv in val.items():
            doc.add_heading(heading(sk),2)
            if isinstance(sv,list):
                for item in sv: doc.add_paragraph(str(item),style='List Bullet')
            else: doc.add_paragraph(str(sv))
    else:
        for para in re.split(r'\n\s*\n',str(val)):
            p=doc.add_paragraph(para.strip());p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if k in ['abstract','problemStatement','proposedSystem','methodology','systemArchitecture','testingStrategy']: doc.add_page_break()

doc.save(docx_out)

# ---------- PDF ----------
styles=getSampleStyleSheet()
styles.add(ParagraphStyle(name='PMTitle',parent=styles['Title'],fontName='Helvetica-Bold',fontSize=23,leading=29,textColor=colors.HexColor('#29266f'),alignment=TA_CENTER,spaceAfter=22))
styles.add(ParagraphStyle(name='PMH1',parent=styles['Heading1'],fontName='Helvetica-Bold',fontSize=16,leading=20,textColor=colors.HexColor('#353087'),spaceBefore=10,spaceAfter=8))
styles.add(ParagraphStyle(name='PMH2',parent=styles['Heading2'],fontName='Helvetica-Bold',fontSize=12,leading=15,textColor=colors.HexColor('#5148b5'),spaceBefore=7,spaceAfter=5))
styles.add(ParagraphStyle(name='PMBody',parent=styles['BodyText'],fontName='Helvetica',fontSize=10.2,leading=15,alignment=TA_JUSTIFY,spaceAfter=8))
styles.add(ParagraphStyle(name='PMCenter',parent=styles['BodyText'],alignment=TA_CENTER,fontSize=11,leading=16))

def footer(canvas,doc):
    canvas.saveState();canvas.setFont('Helvetica',8);canvas.setFillColor(colors.grey);canvas.drawCentredString(A4[0]/2,20,f"PresentMind AI • {data.get('title','Project Synopsis')} • Page {doc.page}");canvas.restoreState()

story=[Spacer(1,1.0*inch),Paragraph(data.get('title','Project Synopsis'),styles['PMTitle']),Paragraph('PROJECT SYNOPSIS',styles['PMCenter']),Spacer(1,.25*inch)]
meta_data=[[Paragraph(f'<b>{a}</b>',styles['PMBody']),Paragraph(words(b),styles['PMBody'])] for a,b in meta]
table=Table(meta_data,colWidths=[1.65*inch,4.7*inch]);table.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.5,colors.HexColor('#d6d6ea')),('BACKGROUND',(0,0),(0,-1),colors.HexColor('#efeffb')),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),8),('RIGHTPADDING',(0,0),(-1,-1),8)]));story+=[table,Spacer(1,.5*inch),Paragraph('Intelligent Presentation & Research Workspace',styles['PMCenter']),PageBreak()]
story += [Paragraph('Declaration',styles['PMH1']),Paragraph(f"I, {data.get('studentName','Vishal Dubey')}, declare that the project synopsis entitled <b>{data.get('title','')}</b> represents an original proposed work prepared for academic and development purposes. The proposed design, objectives, methodology and implementation plan have been described to the best of my understanding. Any external concepts, libraries or research ideas used during implementation will be properly acknowledged in the final report.",styles['PMBody']),Spacer(1,.5*inch),Paragraph('Student Signature: ____________________',styles['PMBody']),Paragraph('Date: ____________________',styles['PMBody']),PageBreak(),Paragraph('Table of Contents',styles['PMH1'])]
for i,k in enumerate(order,1): story.append(Paragraph(f'{i}. {heading(k)}',styles['PMBody']))
story.append(PageBreak())
for k in order:
    val=data.get(k)
    if not val: continue
    story.append(Paragraph(heading(k),styles['PMH1']))
    if isinstance(val,list):
        items=[]
        for item in val:
            if isinstance(item,dict):
                title=item.get('title') or item.get('name') or item.get('phase') or 'Item'; desc=item.get('description') or item.get('details') or item.get('duration') or ''
                txt=f'<b>{title}</b>{": "+str(desc) if desc else ""}'
            else: txt=str(item)
            items.append(ListItem(Paragraph(txt,styles['PMBody']),leftIndent=14))
        story.append(ListFlowable(items,bulletType='bullet',leftIndent=18))
    elif isinstance(val,dict):
        for sk,sv in val.items():
            story.append(Paragraph(heading(sk),styles['PMH2']))
            if isinstance(sv,list): story.append(ListFlowable([ListItem(Paragraph(str(x),styles['PMBody']),leftIndent=14) for x in sv],bulletType='bullet',leftIndent=18))
            else: story.append(Paragraph(str(sv),styles['PMBody']))
    else:
        for para in re.split(r'\n\s*\n',str(val)): story.append(Paragraph(para.strip(),styles['PMBody']))
    if k in ['abstract','problemStatement','proposedSystem','methodology','systemArchitecture','testingStrategy']: story.append(PageBreak())
SimpleDocTemplate(pdf_out,pagesize=A4,rightMargin=45,leftMargin=45,topMargin=45,bottomMargin=38,title=data.get('title','Project Synopsis')).build(story,onFirstPage=footer,onLaterPages=footer)
